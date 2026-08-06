import io

import pytest

from services.file_text import UnsupportedFileError, extract_text


def test_extract_txt_decodes_utf8_and_strips_bom():
    data = "Hello.\nWorld.".encode("utf-8-sig")  # utf-8-sig prepends the BOM itself

    text = extract_text("sample.txt", data)

    assert text == "Hello.\nWorld."


def test_extract_txt_is_case_insensitive_on_extension():
    data = b"Hello."

    text = extract_text("SAMPLE.TXT", data)

    assert text == "Hello."


def test_extract_unsupported_extension_raises():
    with pytest.raises(UnsupportedFileError):
        extract_text("sample.xyz", b"whatever")


def test_extract_no_extension_raises():
    with pytest.raises(UnsupportedFileError):
        extract_text("sample", b"whatever")


def test_extract_docx_reads_paragraph_text():
    from docx import Document

    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("")
    document.add_paragraph("Second paragraph.")
    document.save(buffer)

    text = extract_text("sample.docx", buffer.getvalue())

    assert text == "First paragraph.\n\nSecond paragraph."


def test_extract_pdf_reads_page_text():
    from pypdf import PdfWriter

    buffer = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.write(buffer)

    # A blank page has no extractable text; this should not raise and
    # should return an empty string rather than crashing.
    text = extract_text("sample.pdf", buffer.getvalue())

    assert text == ""
